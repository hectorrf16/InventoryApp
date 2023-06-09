#Signing the code
#Héctor Rodríguez Fusté
#This work is licensed under the Creative Commons Attribution-ShareAlike 4.0 International License.
#To view a copy of this license, visit http://creativecommons.org/licenses/by-sa/4.0/.
#Version: 1 - Created on 24/10/2017
ProgramVersion = "2"
DateVersion = "07/11/2019"

#importing the libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import load_workbook
from openpyxl import Workbook
import shutil
import os
import subprocess
import time
import logging
import comtypes.client
from easygui import *


def covx_to_pdf(infile):
	"""Convert a Word .docx to PDF"""
	wdFormatPDF = 17
	pwd = os.getcwd().replace("\\","\\\\")

	infile = pwd + "\\" + infile
	outfile = infile.replace(".docx",".pdf")
	word = comtypes.client.CreateObject('Word.Application')
	doc = word.Documents.Open(infile)
	doc.SaveAs(outfile, FileFormat=wdFormatPDF)
	doc.Close()
	word.Quit()
def charCounter(string = None, char = " "):
    var = len(string) - len(string.replace(char,""))
    return var

os.makedirs("logs", exist_ok=True)
logfile = "logs\\LogFile " + time.strftime("%A %d %B") + ".log"

WordProgram = "\"C:\\Program Files (x86)\\Microsoft Office\\Office16\\WINWORD.EXE\""

logging.basicConfig(level=logging.DEBUG, filename=logfile)

try:
	# create folders, no error if it already exists
	MonthName= time.strftime("%B")
	MonthNumber= time.strftime("%m")
	YearNumber = time.strftime("%Y")
	DayNumber = time.strftime("%d")
	TodayDate = time.strftime("%d/%m/%Y")
	months = ["Unknown","January","February","March","April","May","June","July","August","September","October","November","December"]
	OutputFiles = "Inventory"
	os.makedirs(OutputFiles, exist_ok=True)
	os.makedirs(OutputFiles + "\\DB", exist_ok=True)
	os.makedirs(OutputFiles + "\\DB" + "\\BK", exist_ok=True)
	os.makedirs(OutputFiles + "\\Reports", exist_ok=True)

    #FIRST WINDOW
	ProgramTitle = "Inventory Creator"
	MainMenuMsg = "Program Version: " + ProgramVersion + "\nDate Version: " + DateVersion + "\n\nWhat would you like to do? Click a button to choose an option"
	MainMenuChoices= ["[1] Create Inventory", "[2] Get Report", "[3] Search", "[4] Exit"]
	MainMenuChoice = "4"

	#INVENTORY WINDOW
	InventoryMenuMsg = "What Item would you like to add?"
	InventoryMenuChoices = ["[1] Add Phone", "[2] Add Workstation", "[3] Add Laptop", "[4] Add Printer", "[5] Add Miscellaneous", "[6] Back"]
	InsInvDateFields = ["Month", "Year"]
	InsInvDateDefaultData = ["MM", "YYYY"]

	#OPTIONS FOR INVENTORY (WS/LT/MISC)
	AddItemInvFileMsg = "How many Items will you add into the new Inventory File?"

	PhonesInvFileFields = ["Brand Name","Product Name","Model","Product Number","Serial Number","IMEI","Date"]
	PhonesDefaultInvData = ["No Data", "No Data", "No Data", "No Data", "No Data","No Data","DD/MM/YYYY"]
	
	WSInvFileFields = ["Brand Name","Product Name","Model","Product Number","Serial Number","Date"]
	WSDefaultInvData = ["No Data", "No Data", "No Data", "No Data", "No Data","DD/MM/YYYY"]
	
	LTInvFileFields = ["Brand Name","Product Name","Model","Product Number","Serial Number","Date"]
	LTDefaultInvData = ["No Data", "No Data", "No Data", "No Data", "No Data","DD/MM/YYYY"]
	
	PRTInvFileFields = ["Brand Name","Product Name","Model","Product Number","Serial Number","Date"]
	PRTDefaultInvData = ["No Data", "No Data", "No Data", "No Data", "No Data","DD/MM/YYYY"]
	
	MicInvFileFields = ["Brand Name","Product Name","Model","Product Number","Serial Number","Date"]
	MicDefaultInvData = ["No Data", "No Data", "No Data", "No Data", "No Data","DD/MM/YYYY"]
	
	SearchItemMenuMsg = "What Item would you like to look for?"
	SearchReportMenuMsg = "What type of Report would you like get?"

	SearchReportMenuChoices = ["[1] All Items all Year", "[2] All Items per Month", "[3] A type of Item per Year", "[4] A type of Item per Month", "[5] Back"]
	SearchItemMenuChoices = ["[1] Phones", "[2] Workstations", "[3] Laptops", "[4] Printers", "[5] Miscellaneous", "[6] Back"]

	InventoryFileName = "Inventory.xlsx"
	BkFileName = "Inventory - " + MonthName + " of " + YearNumber + ".xlsx"
	InventoryFileDir = ".\\Inventory\\DB\\" + InventoryFileName
	BkFileDir = ".\\Inventory\\DB\\BK\\" + BkFileName
	Exit = "No"

	while Exit == "No":

		MainMenuChoice = ((buttonbox(msg=MainMenuMsg, title=ProgramTitle, choices=MainMenuChoices)).split("] ")[0]).replace("[","")
		if MainMenuChoice == "1":
			if os.path.exists(InventoryFileDir) == False:
				docE = load_workbook(".\\Templates\\ExcelTemplate.xlsx") #File where we have the data saved.
				Data = docE["Inventory"] #The name of the data sheet
				Data.cell(row=1,column=2, value=TodayDate)
				Data.cell(row=2,column=2, value=TodayDate)

			else:
				docE = load_workbook(InventoryFileDir) #File where we have the data saved.
				Data = docE["Inventory"] #The name of the data sheet
				if Data.cell(row=2,column=2).value == None:
					Data.cell(row=2,column=2, value=TodayDate)
				MLastBkDate = (Data.cell(row=2,column=2).value).split("/")[1]
				LastBkDate = Data.cell(row=2,column=2).value
				DLastBkDate = (Data.cell(row=2,column=2).value).split("/")[0]

				if MonthNumber > (LastBkDate): 
					Data.cell(row=2,column=2, value=TodayDate)
					shutil.copy(InventoryFileDir, BkFileDir)
				elif MonthNumber == MLastBkDate and DayNumber > DLastBkDate:
					Data.cell(row=2,column=2, value=TodayDate)
					shutil.copy(InventoryFileDir, BkFileDir)
				else:
					Data.cell(row=2,column=2, value=TodayDate)

			InventoryMenuChoice = ((buttonbox(msg=InventoryMenuMsg, title=ProgramTitle, choices=InventoryMenuChoices)).split("] ")[0]).replace("[","")
			if InventoryMenuChoice == "1":
				Data = docE["Phones"] #The name of the data sheet.
			elif InventoryMenuChoice == "2":
				Data = docE["Workstations"] #The name of the data sheet.
			elif InventoryMenuChoice == "3":
				Data = docE["Laptops"] #The name of the data sheet.
			elif InventoryMenuChoice == "4":
				Data = docE["Printers"] #The name of the data sheet.
			elif InventoryMenuChoice == "5":
				Data = docE["Miscellaneous"] #The name of the data sheet.
			elif InventoryMenuChoice == "6":
				break
			NumItems = int(enterbox(msg=AddItemInvFileMsg, title=ProgramTitle, default="1"))
			for x in range(2,Data.max_row + NumItems+1):

				if Data.cell(row=x,column=1).value != None:
					next
				else:
					if InventoryMenuChoice == "1":
						PhonesDefaultInvData[6] = TodayDate
						AddDataToInvFile = multenterbox(msg="Enter the Data", title=ProgramTitle, fields=PhonesInvFileFields, values=PhonesDefaultInvData)
						AddDataToInvFile[2] = (AddDataToInvFile[2]).upper()
						AddDataToInvFile[3] = (AddDataToInvFile[3]).upper()
						AddDataToInvFile[4] = (AddDataToInvFile[4]).upper()
						AddDataToInvFile[5] = str(AddDataToInvFile[5]).replace(" ","")
						Data.cell(row=x,column=1, value=AddDataToInvFile[0])
						Data.cell(row=x,column=2, value=AddDataToInvFile[1])
						Data.cell(row=x,column=3, value=AddDataToInvFile[2])
						Data.cell(row=x,column=4, value=AddDataToInvFile[3])
						Data.cell(row=x,column=5, value=AddDataToInvFile[4])
						Data.cell(row=x,column=6, value=AddDataToInvFile[5])
						Data.cell(row=x,column=7, value=AddDataToInvFile[6])
					elif InventoryMenuChoice == "2":
						WSDefaultInvData[5] = TodayDate
						AddDataToInvFile = multenterbox(msg="Enter the Data", title=ProgramTitle, fields=WSInvFileFields, values=WSDefaultInvData)
						AddDataToInvFile[2] = (AddDataToInvFile[2]).upper()
						AddDataToInvFile[3] = (AddDataToInvFile[3]).upper()
						AddDataToInvFile[4] = (AddDataToInvFile[4]).upper()
						Data.cell(row=x,column=1, value=AddDataToInvFile[0])
						Data.cell(row=x,column=2, value=AddDataToInvFile[1])
						Data.cell(row=x,column=3, value=AddDataToInvFile[2])
						Data.cell(row=x,column=4, value=AddDataToInvFile[3])
						Data.cell(row=x,column=5, value=AddDataToInvFile[4])
						Data.cell(row=x,column=6, value=AddDataToInvFile[5])
					elif InventoryMenuChoice == "3":
						LTDefaultInvData[5] = TodayDate
						AddDataToInvFile = multenterbox(msg="Enter the Data", title=ProgramTitle, fields=LTInvFileFields, values=LTDefaultInvData)
						AddDataToInvFile[2] = (AddDataToInvFile[2]).upper()
						AddDataToInvFile[3] = (AddDataToInvFile[3]).upper()
						AddDataToInvFile[4] = (AddDataToInvFile[4]).upper()
						Data.cell(row=x,column=1, value=AddDataToInvFile[0])
						Data.cell(row=x,column=2, value=AddDataToInvFile[1])
						Data.cell(row=x,column=3, value=AddDataToInvFile[2])
						Data.cell(row=x,column=4, value=AddDataToInvFile[3])
						Data.cell(row=x,column=5, value=AddDataToInvFile[4])
						Data.cell(row=x,column=6, value=AddDataToInvFile[5])
					elif InventoryMenuChoice == "4":
						PRTDefaultInvData[5] = TodayDate
						AddDataToInvFile = multenterbox(msg="Enter the Data", title=ProgramTitle, fields=PRTInvFileFields, values=PRTDefaultInvData)
						AddDataToInvFile[2] = (AddDataToInvFile[2]).upper()
						AddDataToInvFile[3] = (AddDataToInvFile[3]).upper()
						AddDataToInvFile[4] = (AddDataToInvFile[4]).upper()
						Data.cell(row=x,column=1, value=AddDataToInvFile[0])
						Data.cell(row=x,column=2, value=AddDataToInvFile[1])
						Data.cell(row=x,column=3, value=AddDataToInvFile[2])
						Data.cell(row=x,column=4, value=AddDataToInvFile[3])
						Data.cell(row=x,column=5, value=AddDataToInvFile[4])
						Data.cell(row=x,column=6, value=AddDataToInvFile[5])
					elif InventoryMenuChoice == "5":
						MicDefaultInvData[5] = TodayDate
						AddDataToInvFile = multenterbox(msg="Enter the Data", title=ProgramTitle, fields=MicInvFileFields, values=MicDefaultInvData)
						AddDataToInvFile[2] = (AddDataToInvFile[2]).upper()
						AddDataToInvFile[3] = (AddDataToInvFile[3]).upper()
						AddDataToInvFile[4] = (AddDataToInvFile[4]).upper()
						Data.cell(row=x,column=1, value=AddDataToInvFile[0])
						Data.cell(row=x,column=2, value=AddDataToInvFile[1])
						Data.cell(row=x,column=3, value=AddDataToInvFile[2])
						Data.cell(row=x,column=4, value=AddDataToInvFile[3])
						Data.cell(row=x,column=5, value=AddDataToInvFile[4])
						Data.cell(row=x,column=6, value=AddDataToInvFile[5])
			else:
				docE.save(InventoryFileDir)
		elif MainMenuChoice == "2":
			SearchReportMenuChoice = ((buttonbox(msg=SearchReportMenuMsg, title=ProgramTitle, choices=SearchReportMenuChoices)).split("] ")[0]).replace("[","")
			
			if SearchReportMenuChoice == "1":
				DateInvFile = enterbox(msg="Enter the Year number", title=ProgramTitle, value="YYYY")
				InventoryReportName = "Inventory - " + DateInvFileMonthName + " of " + DateInvFileYearNumber + ".docx"
				InventoryReportModifiedName = "Inventory - " + DateInvFileMonthName + " of " + DateInvFileYearNumber + " - Modified" + ".docx"
				InventoryReportDir = ".\\Inventory\\Reports\\" + DateInvFileYearNumber + "\\" + DateInvFileMonthName + "\\" + InventoryReportName
				InventoryReportModifiedDir = ".\\Inventory\\Reports\\" + DateInvFileYearNumber + "\\" + DateInvFileMonthName + "\\" + InventoryReportModifiedName

				SkippedLines = 0
				InvReportFile = ".\\Templates\\WordTemplate.docx"
				docW = Document(InvReportFile)

				docW._body.clear_content()
				HeaderTable = docW.add_table(rows=1,cols=1)
				HeaderTable.rows[0].cells[0].text = "Inventory for " +  DateInvFileMonthName + " of " + DateInvFileYearNumber
				HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
				HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)

				if os.path.exists(InventoryFileDir) == False:
					break
				else:
					docE = load_workbook(InventoryFileDir)
					for x in range(1,5 + 1):
						if x == 1:
							Data = docE["Phone"]
						elif x == 2:
							Data = docE["Workstations"]
						elif x == 3:
							Data = docE["Laptops"]
						elif x == 4:
							Data = docE["Printers"]
						elif x == 5:
							Data = docE["Miscellaneous"]

						IsThereAnEntry = False
						for y in range(2,(Data.max_row) + 1):
							if (Data.cell(row=y, column=7).value).split("/")[2] == None:
								SkippedLines+=1
								next
							elif (Data.cell(row=y, column=7).value).split("/")[2] != DateInvFile:
								next
							else:
								if x == 1:
									Templatetable1 = docW.add_table(rows=1, cols=7, style="Table Grid")

									Templatetable1.rows[0].cells[0].text = Data.cell(row=1,column=1).value
									Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[1].text = Data.cell(row=1,column=2).value
									Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[2].text = Data.cell(row=1,column=3).value
									Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[3].text = Data.cell(row=1,column=4).value
									Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[4].text = Data.cell(row=1,column=5).value
									Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[5].text = Data.cell(row=1,column=6).value
									Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[6].text = Data.cell(row=1,column=7).value
									Templatetable1.rows[0].cells[6].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[6].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

								else:
									Templatetable1 = docW.add_table(rows=1, cols=6, style="Table Grid")

									Templatetable1.rows[0].cells[0].text = Data.cell(row=1,column=1).value
									Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[1].text = Data.cell(row=1,column=2).value
									Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[2].text = Data.cell(row=1,column=3).value
									Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[3].text = Data.cell(row=1,column=4).value
									Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[4].text = Data.cell(row=1,column=5).value
									Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
									Templatetable1.rows[0].cells[5].text = Data.cell(row=1,column=6).value
									Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].bold = True
									Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(9)
									Templatetable1.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

			elif SearchReportMenuChoice == "2":
			elif SearchReportMenuChoice == "3":
			elif SearchReportMenuChoice == "4":
			elif SearchReportMenuChoice == "5":
				exit()
			SearchItemMenuChoice = ((buttonbox(msg=SearchItemMenuMsg, title=ProgramTitle, choices=SearchItemMenuChoices)).split("] ")[0]).replace("[","")
			if os.path.exists(InventoryFileDir) == False:
				break
			else:
				docE = load_workbook(InventoryFileDir)
				if SearchItemMenuChoice == "1":
					Data = docE["Phones"] #The name of the data sheet.
				elif SearchItemMenuChoice == "2":
					Data = docE["Workstations"] #The name of the data sheet.
				elif SearchItemMenuChoice == "3":
					Data = docE["Laptops"] #The name of the data sheet.
				elif SearchItemMenuChoice == "4":
					Data = docE["Printers"] #The name of the data sheet.
				elif SearchItemMenuChoice == "5":
					Data = docE["Miscellaneous"] #The name of the data sheet.
				elif SearchItemMenuChoice == "6":
					break

			DateInvFile = multenterbox(msg="Enter the Data", title=ProgramTitle, fields=InsInvDateFields, values=InsInvDateDefaultData)
			DateInvFileMonthName = months[int(DateInvFile[0])]
			DateInvFileYearNumber = DateInvFile[1]

			IsThereAnEntry = False
			for y in range(2,(Data.max_row) + 1):
				if SearchItemMenuChoice == "1":
					if (Data.cell(row=y, column=7).value).split("/")[1] != DateInvFile[0]:
						next
					else:
						IsThereAnEntry = True
				elif SearchItemMenuChoice != "1":
					if (Data.cell(row=y, column=6).value).split("/")[1] != DateInvFile[0]:
						next
					else:
						IsThereAnEntry = True
			else:
				if IsThereAnEntry == False:
					msgbox(msg="No Items detected on this Date")
					exit()

			InventoryReportName = " Inventory - " + DateInvFileMonthName + " of " + DateInvFileYearNumber + ".docx"
			InventoryReportModifiedName = " Inventory - " + DateInvFileMonthName + " of " + DateInvFileYearNumber + " - Modified" + ".docx"
			InventoryReportDir = ".\\Inventory\\Reports\\" + DateInvFileYearNumber + "\\" + DateInvFileMonthName + "\\" + InventoryReportName
			InventoryReportModifiedDir = ".\\Inventory\\Reports\\" + DateInvFileYearNumber + "\\" + DateInvFileMonthName + "\\" + InventoryReportModifiedName

			SkippedLines = 0
			InvReportFile = ".\\Templates\\WordTemplate.docx"
			docW = Document(InvReportFile)

			docW._body.clear_content()
			HeaderTable = docW.add_table(rows=1,cols=1)
			HeaderTable.rows[0].cells[0].text = "Inventory for " +  DateInvFileMonthName + " of " + DateInvFileYearNumber
			HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].bold = True
			HeaderTable.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
			if SearchItemMenuChoice == "1":
				Templatetable1 = docW.add_table(rows=1, cols=7, style="Table Grid")

				Templatetable1.rows[0].cells[0].text = Data.cell(row=1,column=1).value
				Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[1].text = Data.cell(row=1,column=2).value
				Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[2].text = Data.cell(row=1,column=3).value
				Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[3].text = Data.cell(row=1,column=4).value
				Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[4].text = Data.cell(row=1,column=5).value
				Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[5].text = Data.cell(row=1,column=6).value
				Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[6].text = Data.cell(row=1,column=7).value
				Templatetable1.rows[0].cells[6].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[6].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
			else:
				Templatetable1 = docW.add_table(rows=1, cols=6, style="Table Grid")

				Templatetable1.rows[0].cells[0].text = Data.cell(row=1,column=1).value
				Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[1].text = Data.cell(row=1,column=2).value
				Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[2].text = Data.cell(row=1,column=3).value
				Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[3].text = Data.cell(row=1,column=4).value
				Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[4].text = Data.cell(row=1,column=5).value
				Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				Templatetable1.rows[0].cells[5].text = Data.cell(row=1,column=6).value
				Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].bold = True
				Templatetable1.rows[0].cells[5].paragraphs[0].runs[0].font.size = Pt(9)
				Templatetable1.rows[0].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

			for x in range(2,(Data.max_row) + 1):
				if SearchItemMenuChoice != "1":
					Date = Data.cell(row=x,column=6).value
					if Date.split("/")[2] != DateInvFileYearNumber or Date.split("/")[1] != DateInvFile[0]:
						next
						SkippedLines = SkippedLines + 1
					else:
						BrandName = Data.cell(row=x,column=1).value
						ProductName = Data.cell(row=x,column=2).value
						Model = (Data.cell(row=x,column=3).value).upper()
						ProductNumber = Data.cell(row=x,column=4).value
						if ProductNumber == "No Data":
							ProductNumber = Data.cell(row=x,column=4).value
						else:
							ProductNumber = (Data.cell(row=x,column=4).value).upper()
						SerialNumber = (Data.cell(row=x,column=5).value).upper()

						Templatetable1.add_row()
						Templatetable1.rows[(x-1)-SkippedLines].cells[0].text = BrandName
						Templatetable1.rows[(x-1)-SkippedLines].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[1].text = ProductName
						Templatetable1.rows[(x-1)-SkippedLines].cells[1].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[2].text = Model
						Templatetable1.rows[(x-1)-SkippedLines].cells[2].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[3].text = ProductNumber
						Templatetable1.rows[(x-1)-SkippedLines].cells[3].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[4].text = SerialNumber
						Templatetable1.rows[(x-1)-SkippedLines].cells[4].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[5].text = Date
						Templatetable1.rows[(x-1)-SkippedLines].cells[5].paragraphs[0].runs[0].font.size = Pt(8)
				else:
					Date = Data.cell(row=x,column=7).value
					if Date.split("/")[2] != DateInvFileYearNumber or Date.split("/")[1] != DateInvFile[0]:
						next
						SkippedLines = SkippedLines + 1
					else:
						BrandName = Data.cell(row=x,column=1).value
						ProductName = Data.cell(row=x,column=2).value
						Model = (Data.cell(row=x,column=3).value).upper()
						ProductNumber = Data.cell(row=x,column=4).value
						if ProductNumber == "No Data":
							ProductNumber = Data.cell(row=x,column=4).value
						else:
							ProductNumber = (Data.cell(row=x,column=4).value).upper()
						SerialNumber = (Data.cell(row=x,column=5).value).upper()
						IMEINumber = Data.cell(row=x,column=6).value

						Templatetable1.add_row()
						Templatetable1.rows[(x-1)-SkippedLines].cells[0].text = BrandName
						Templatetable1.rows[(x-1)-SkippedLines].cells[0].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[1].text = ProductName
						Templatetable1.rows[(x-1)-SkippedLines].cells[1].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[2].text = Model
						Templatetable1.rows[(x-1)-SkippedLines].cells[2].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[3].text = ProductNumber
						Templatetable1.rows[(x-1)-SkippedLines].cells[3].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[4].text = SerialNumber
						Templatetable1.rows[(x-1)-SkippedLines].cells[4].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[5].text = IMEINumber
						Templatetable1.rows[(x-1)-SkippedLines].cells[5].paragraphs[0].runs[0].font.size = Pt(8)
						Templatetable1.rows[(x-1)-SkippedLines].cells[6].text = Date
						Templatetable1.rows[(x-1)-SkippedLines].cells[6].paragraphs[0].runs[0].font.size = Pt(8)
			else:
				os.makedirs(OutputFiles + "\\Reports\\" + DateInvFileYearNumber, exist_ok=True)
				os.makedirs(OutputFiles + "\\Reports\\" + DateInvFileYearNumber + "\\" + DateInvFileMonthName, exist_ok=True)

				if os.path.exists(InventoryReportDir) == True and os.path.isfile(InventoryReportDir) == True:
					ReplaceOption = boolbox(msg="The Report you want to create already exists. Do you want to create a Modified File ? If you click \"No\", the old Report will be overwritten")
					if ReplaceOption == True:
						docW.save(InventoryReportModifiedDir)
						covx_to_pdf(InventoryReportModifiedDir)
						subprocess.call(WordProgram + "\"" + InventoryReportModifiedDir + "\"")
					else:
						docW.save(InventoryReportDir)
						covx_to_pdf(InventoryReportDir)
						
						subprocess.call(WordProgram + "\"" + InventoryReportDir + "\"")
				else:
					docW.save(InventoryReportDir)
					covx_to_pdf(InventoryReportDir)

					subprocess.call(WordProgram + "\"" + InventoryReportDir + "\"")
		elif MainMenuChoice == "3":
			SearchItemMenuChoice = ((buttonbox(msg=SearchItemMenuMsg, title=ProgramTitle, choices=SearchItemMenuChoices)).split("] ")[0]).replace("[","")
			docE = load_workbook(InventoryFileDir)
			if SearchItemMenuChoice == "1":
				Data = docE["Phones"] #The name of the data sheet.
				SearchItem = enterbox(title=ProgramTitle, msg="Enter the Item \"Serial Number\"")
			elif SearchItemMenuChoice == "2":
				Data = docE["Workstations"] #The name of the data sheet.
				SearchItem = enterbox(title=ProgramTitle, msg="Enter the Item \"Serial Number\"")
			elif SearchItemMenuChoice == "3":
				Data = docE["Laptops"] #The name of the data sheet.
				SearchItem = enterbox(title=ProgramTitle, msg="Enter the Item \"Serial Number\"")
			elif SearchItemMenuChoice == "4":
				Data = docE["Printers"] #The name of the data sheet.
				SearchItem = enterbox(title=ProgramTitle, msg="Enter the Item \"Serial Number\"")
			elif SearchItemMenuChoice == "5":
				Data = docE["Miscellaneous"] #The name of the data sheet.
				SearchItem = enterbox(title=ProgramTitle, msg="Enter the Item \"Serial Number\"")
			elif SearchItemMenuChoice == "6":
				break

			SearchItem = str(SearchItem).upper()
			for x in range(2,Data.max_row+1):
				if Data.cell(row=x,column=5).value == SearchItem:
					if SearchItemMenuChoice == "1":
						ccbox(title=ProgramTitle, msg="Brand Name: " + Data.cell(row=x,column=1).value + "\nProduct Name: " + Data.cell(row=x,column=2).value + "\nModel: " + Data.cell(row=x,column=3).value + "\nProduct Number: " + Data.cell(row=x,column=4).value + "\nSerial Number: " + Data.cell(row=x,column=5).value + "\nIMEI: " + Data.cell(row=x,column=6).value + "\nDate: " + Data.cell(row=x,column=7).value)
					elif SearchItemMenuChoice == "2":
						ccbox(title=ProgramTitle, msg="Brand Name: " + Data.cell(row=x,column=1).value + "\nProduct Name: " + Data.cell(row=x,column=2).value + "\nModel: " + Data.cell(row=x,column=3).value + "\nProduct Number: " + Data.cell(row=x,column=4).value + "\nSerial Number: " + Data.cell(row=x,column=5).value + "\nDate: " + Data.cell(row=x,column=6).value)
					elif SearchItemMenuChoice == "3":
						ccbox(title=ProgramTitle, msg="Brand Name: " + Data.cell(row=x,column=1).value + "\nProduct Name: " + Data.cell(row=x,column=2).value + "\nModel: " + Data.cell(row=x,column=3).value + "\nProduct Number: " + Data.cell(row=x,column=4).value + "\nSerial Number: " + Data.cell(row=x,column=5).value + "\nDate: " + Data.cell(row=x,column=6).value)
					elif SearchItemMenuChoice == "4":
						ccbox(title=ProgramTitle, msg="Brand Name: " + Data.cell(row=x,column=1).value + "\nProduct Name: " + Data.cell(row=x,column=2).value + "\nModel: " + Data.cell(row=x,column=3).value + "\nProduct Number: " + Data.cell(row=x,column=4).value + "\nSerial Number: " + Data.cell(row=x,column=5).value + "\nDate: " + Data.cell(row=x,column=6).value)
					elif SearchItemMenuChoice == "5":	
						ccbox(title=ProgramTitle, msg="Brand Name: " + Data.cell(row=x,column=1).value + "\nProduct Name: " + Data.cell(row=x,column=2).value + "\nModel: " + Data.cell(row=x,column=3).value + "\nProduct Number: " + Data.cell(row=x,column=4).value + "\nSerial Number: " + Data.cell(row=x,column=5).value + "\nDate: " + Data.cell(row=x,column=6).value)
		elif MainMenuChoice == "4":
			Exit = "Yes"
			shutil.copy(InventoryFileDir, BkFileDir)
except:
	logging.exception("Errors at " + time.strftime("%H:%M:%S") + "\n")